# -*- coding: utf-8 -*-
"""
bid_generator/doc_builder.py
=============================
Word 文档构建器：
  - 读取应标方案模板（Markdown）
  - 按模板章节结构创建 Word 文档
  - 调用生成器填充各章节内容
  - 支持标题、段落、列表、表格等格式
"""

import logging
import re
from pathlib import Path
from typing import Dict, List, Any, Optional

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from config import TEMPLATE_PATH, logger

# ========== 模板解析 ==========

def parse_template(template_path: str = None) -> List[Dict[str, Any]]:
    """
    解析 Markdown 格式的模板文件，返回章节结构列表。

    返回格式:
    [{
        "level": 1/2/3,
        "title": "章节标题",
        "desc": "描述/子标题（可选）",
        "is_leaf": True/False,   # 是否为叶子章节（需填充内容）
    }, ...]
    """
    path = Path(template_path or TEMPLATE_PATH)
    if not path.exists():
        raise FileNotFoundError(f"模板文件不存在: {path}")

    content = path.read_text(encoding="utf-8")
    chapters = []

    # 解析 Markdown 标题行
    for line in content.splitlines():
        line = line.strip()
        if not line:
            continue
        m = re.match(r"^(#{1,3})\s+(.*)", line)
        if not m:
            continue
        level = len(m.group(1))
        title = m.group(2).strip()
        # 去掉 Markdown 链接语法
        title = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", title)
        chapters.append({
            "level":   level,
            "title":   title,
            "desc":    "",
            "is_leaf": level == 3,
        })

    logger.info(f"[DocBuilder] 模板解析完成，共 {len(chapters)} 个章节")
    return chapters


def get_leaf_chapters(chapters: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """提取所有叶子章节（H3）"""
    return [c for c in chapters if c.get("is_leaf", False)]


# ========== Word 文档样式工具 ==========

def set_heading_style(paragraph, level: int):
    """设置 Word 标题样式"""
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.font.size = {1: Pt(20), 2: Pt(16), 3: Pt(14)}.get(level, Pt(14))
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_paragraph_with_style(doc: Document, text: str, bold: bool = False, size: int = 11):
    """添加普通段落"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.5
    return para


def add_bullet_list(doc: Document, items: List[str]):
    """添加项目符号列表"""
    for item in items:
        para = doc.add_paragraph(style="List Bullet")
        para.add_run(item)
        para.paragraph_format.space_after = Pt(2)


def add_table_from_text(doc: Document, text: str):
    """
    尝试从文本中解析简单表格（用 | 分隔）。
    无法解析时直接添加为普通段落。
    """
    lines = [l.strip() for l in text.splitlines() if "|" in l and l.strip()]
    if not lines:
        add_paragraph_with_style(doc, text)
        return

    # 简单表格：| 列1 | 列2 | ...
    rows_data = []
    for line in lines:
        cells = [c.strip() for c in line.split("|") if c.strip()]
        rows_data.append(cells)

    if not rows_data:
        add_paragraph_with_style(doc, text)
        return

    # 最多取前10行、8列
    rows_data = rows_data[:10]
    cols = max(len(r) for r in rows_data)
    rows_data = [r + [""] * (cols - len(r)) for r in rows_data]

    table = doc.add_table(rows=len(rows_data), cols=cols)
    table.style = "Table Grid"
    for i, row_data in enumerate(rows_data):
        for j, cell_text in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = cell_text
            if i == 0:
                # 表头加粗
                for run in cell.paragraphs[0].runs:
                    run.bold = True
    doc.add_paragraph()  # 空行


# ========== DocBuilder 主类 ==========

class BidDocBuilder:
    """
    标书 Word 文档构建器：
      - 读取模板章节结构
      - 按层级构建 Word 文档（自动生成章/节/子节标题）
      - 向叶子章节填入生成内容
    """

    def __init__(self, template_path: str = None):
        self.template_path = template_path or str(TEMPLATE_PATH)
        self.chapters: List[Dict[str, Any]] = []
        self._doc: Optional[Document] = None

    def load_template(self) -> List[Dict[str, Any]]:
        """加载并解析模板"""
        self.chapters = parse_template(self.template_path)
        return self.chapters

    def build(
        self,
        chapter_contents: Dict[str, str],
        output_path: str,
        project_name: str = "投标方案",
    ):
        """
        构建完整的 Word 文档。

        参数:
          chapter_contents: {章节标题: 生成内容}
          output_path:     输出 .docx 路径
        """
        self.load_template()
        doc = Document()
        self._doc = doc

        # 页面设置
        section = doc.sections[0]
        section.page_width  = Inches(21)   # A4 宽
        section.page_height = Inches(29.7)  # A4 高
        section.left_margin   = Inches(1.8)
        section.right_margin  = Inches(1.8)
        section.top_margin    = Inches(1.5)
        section.bottom_margin = Inches(1.5)

        # 封面（自动生成）
        self._add_cover(doc, project_name)

        # 按层级构建章节
        prev_level = 0
        for ch in self.chapters:
            level = ch["level"]
            title = ch["title"]
            content = chapter_contents.get(title, "")

            # 调整标题层级（Word 中 Heading 1=章, Heading 2=节, Heading 3=子节）
            heading_style = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}.get(level, "Heading 3")

            # 添加标题
            h = doc.add_heading(title, level=level)
            h.style = doc.styles[heading_style]

            # 如果是叶子章节，填入生成内容
            if ch.get("is_leaf") and content:
                self._fill_content(doc, content)
            else:
                # 非叶子章节，空行分隔
                pass

            prev_level = level

        # 保存
        doc.save(output_path)
        logger.info(f"[DocBuilder] Word 文档已保存: {output_path}")

    def _add_cover(self, doc: Document, project_name: str):
        """添加封面"""
        # 标题
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run("\n\n\n\n")
        run.font.size = Pt(28)
        run.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)

        title_para2 = doc.add_paragraph()
        title_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = title_para2.add_run(project_name)
        run2.font.size = Pt(28)
        run2.bold = True
        run2.font.color.rgb = RGBColor(0, 51, 102)

        doc.add_page_break()

    def _fill_content(self, doc: Document, content: str):
        """
        将生成的内容填入文档，智能识别：
          - 列表项（- 或 * 或 编号）
          - 表格（| 分隔）
          - 普通段落
        """
        lines = content.splitlines()
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            i += 1
            if not line:
                continue

            # 跳过 Markdown 标记残留
            if line.startswith("#") or line.startswith("---"):
                continue

            # 列表项（-、*、、或 1. 2. 编号）
            if re.match(r"^[-*·]\s+", line) or re.match(r"^[①②③④⑤]\s+", line) or re.match(r"^\d+[.)]\s+", line):
                # 收集连续列表
                bullet_lines = []
                while i < len(lines) and (re.match(r"^[-*·]\s+", lines[i].strip()) or
                                           re.match(r"^[①②③④⑤]\s+", lines[i].strip()) or
                                           re.match(r"^\d+[.)]\s+", lines[i].strip())):
                    bullet_lines.append(lines[i].strip())
                    i += 1
                self._add_bullet_block(doc, bullet_lines)

            # 表格行（包含 |）
            elif "|" in line and line.count("|") >= 2:
                table_lines = []
                while i < len(lines) and "|" in lines[i] and lines[i].strip().count("|") >= 2:
                    table_lines.append(lines[i].strip())
                    i += 1
                add_table_from_text(doc, "\n".join(table_lines))

            # 普通段落
            else:
                # 合并连续的非列表段落
                para_lines = [line]
                while i < len(lines) and lines[i].strip() and not lines[i].strip().startswith("#") and "|" not in lines[i].strip()[:5]:
                    next_line = lines[i].strip()
                    if not (re.match(r"^[-*·]\s+", next_line) or re.match(r"^\d+[.)]\s+", next_line)):
                        para_lines.append(next_line)
                        i += 1
                    else:
                        break
                para_text = " ".join(para_lines)
                add_paragraph_with_style(doc, para_text)

    def _add_bullet_block(self, doc: Document, bullet_lines: List[str]):
        """添加一个列表块"""
        for bl in bullet_lines:
            # 去掉列表前缀
            text = re.sub(r"^[-*·]\s+", "", bl)
            text = re.sub(r"^[①②③④⑤]\s+", "", text)
            text = re.sub(r"^\d+[.)]\s+", "", text)
            para = doc.add_paragraph(style="List Bullet")
            para.add_run(text)
            para.paragraph_format.space_after = Pt(3)
