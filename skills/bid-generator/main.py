# -*- coding: utf-8 -*-
"""
bid_generator/main.py
======================
应标方案自动生成系统 - 主入口

核心流程：
  招标需求 → 解析关键词 → 向量检索(LanceDB) → AI生成 → Word文档
"""

import logging
import sys
import time
from pathlib import Path
from typing import Union, List, Dict, Any, Optional

from config import (
    TEMPLATE_PATH, WORKSPACE, BID_DOCS_DIR,
    logger, TOP_K, GENERATE_MODEL,
)
from knowledge_base import BidKnowledgeBase
from retriever import BidRetriever, CHAPTER_FILTERS
from generator import BidGenerator
from doc_builder import BidDocBuilder

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s - %(message)s",
    datefmt="%H:%M:%S",
)


# =============================================================================
# 模板解析：生成章节列表（供检索和生成使用）
# =============================================================================

def load_template_chapters() -> List[Dict[str, str]]:
    """从模板文件加载章节列表（用于遍历生成）"""
    from doc_builder import parse_template

    chapters = parse_template(str(TEMPLATE_PATH))
    # 只取 H3 叶子章节（最细粒度的内容单元）
    leaf_chapters = []
    for ch in chapters:
        if ch.get("is_leaf"):
            leaf_chapters.append({
                "title": ch["title"],
                "desc":  ch.get("desc", ""),
                "level": ch["level"],
            })
    return leaf_chapters


# =============================================================================
# 主生成流程
# =============================================================================

def main(
    requirement: Union[str, Path],
    output_path: str = None,
    build_kb: bool = True,
    model: str = None,
    max_chapters: int = 0,  # 0 = 生成全部
) -> str:
    """
    应标方案生成主函数。

    参数:
      requirement: 招标需求（字符串或 .txt/.docx 文件路径）
      output_path: 输出 Word 路径，默认 test_output.docx
      build_kb:    是否重建知识库（False 则复用现有索引）
      model:       覆盖默认生成模型
      max_chapters: 最大生成章节数（0=全部）

    返回:
      输出文件路径
    """
    t0 = time.time()

    # -------------------------------------------------------------------------
    # Step 0: 读取招标需求
    # -------------------------------------------------------------------------
    if isinstance(requirement, (str, Path)):
        req_path = Path(requirement)
        if req_path.exists():
            if req_path.suffix == ".docx":
                from docx import Document
                doc = Document(str(req_path))
                requirement_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            else:
                requirement_text = req_path.read_text(encoding="utf-8")
            logger.info(f"[Step 0] 从文件读取招标需求: {req_path.name}（{len(requirement_text)} 字）")
        else:
            requirement_text = str(requirement)
            logger.info(f"[Step 0] 直接使用传入字符串（{len(requirement_text)} 字）")
    else:
        requirement_text = str(requirement)

    # -------------------------------------------------------------------------
    # Step 1: 初始化 / 更新知识库
    # -------------------------------------------------------------------------
    kb = BidKnowledgeBase()
    if build_kb:
        logger.info("[Step 1] 开始构建 / 更新知识库...")
        index_results = kb.index_all(force=False)
        total_new = sum(v for v in index_results.values() if v > 0)
        logger.info(f"[Step 1] 索引完成，共新增 {total_new} 个 chunk，现有 {kb.count()} 个")
    else:
        logger.info(f"[Step 1] 跳过建库，当前知识库有 {kb.count()} 个 chunk")

    # -------------------------------------------------------------------------
    # Step 2: 初始化检索器和生成器
    # -------------------------------------------------------------------------
    retriever = BidRetriever(kb=kb)
    generator = BidGenerator(model=model or GENERATE_MODEL)

    # 检查模型是否可用
    if not generator.check_model_available():
        logger.error(f"[Generator] 生成模型不可用: {model or GENERATE_MODEL}，请先运行: ollama pull {model or GENERATE_MODEL}")
        sys.exit(1)

    # -------------------------------------------------------------------------
    # Step 3: 按模板章节逐章生成
    # -------------------------------------------------------------------------
    logger.info("[Step 3] 开始按章节生成内容...")
    template_chapters = load_template_chapters()
    if max_chapters > 0:
        template_chapters = template_chapters[:max_chapters]

    logger.info(f"[Step 3] 共需生成 {len(template_chapters)} 个章节")

    # 用于传给生成器的上下文映射
    context_map: Dict[str, List[Dict[str, Any]]] = {}

    generated_contents: Dict[str, str] = {}

    for i, ch in enumerate(template_chapters, 1):
        title = ch["title"]
        logger.info(f"[Step 3] [{i}/{len(template_chapters)}] 处理章节: {title}")

        # 检索相关段落：将章节标题/描述作为检索引导，同时结合招标需求
        search_query = f"{title} {ch['desc']} {requirement_text[:300]}"
        chunks = retriever.retrieve(
            search_query,
            top_k=TOP_K,
            chapter_filter=title[:6],  # 用标题前6字做过滤
        )
        context_map[title] = chunks

        # 生成内容
        content = generator.generate_chapter(
            chapter_title=title,
            chapter_desc=ch["desc"],
            requirement=requirement_text,
            context_chunks=chunks,
            model=model,
        )
        generated_contents[title] = content

        # 每个章节后短暂暂停
        time.sleep(0.3)

    # -------------------------------------------------------------------------
    # Step 4: 构建 Word 文档
    # -------------------------------------------------------------------------
    output_path = output_path or str(WORKSPACE / "test_output.docx")
    logger.info(f"[Step 4] 构建 Word 文档: {output_path}")

    builder = BidDocBuilder()
    builder.build(
        chapter_contents=generated_contents,
        output_path=output_path,
        project_name="投标方案",
    )

    elapsed = time.time() - t0
    logger.info(f"✅ 完成！耗时 {elapsed:.1f}s，输出: {output_path}")
    return output_path


# =============================================================================
# 入口
# =============================================================================

if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="应标方案自动生成系统")
    parser.add_argument("--requirement", "-r", type=str, default=None,
                        help="招标需求文本或 .txt/.docx 文件路径")
    parser.add_argument("--output", "-o", type=str, default=None,
                        help="输出 Word 路径（默认: test_output.docx）")
    parser.add_argument("--no-build-kb", action="store_true",
                        help="跳过知识库构建，复用现有索引")
    parser.add_argument("--model", "-m", type=str, default=None,
                        help="Ollama 生成模型（默认: qwen2.5:7b）")
    parser.add_argument("--max-chapters", type=int, default=0,
                        help="最多生成章节数（0=全部，默认0）")
    args = parser.parse_args()

    # 默认测试需求：康方隆跃方案前200字
    if args.requirement is None:
        from docx import Document
        doc = Document("/Users/visionsl/Documents/资料/标书方案/康方隆跃新药研发与产业化基地建设项目（设备管理系统）-技术标书.docx")
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        test_req = "\n".join(paragraphs[:15])[:500]
        logger.info(f"[Test] 使用康方隆跃方案前500字作为测试需求:\n{test_req[:200]}...")
        requirement_text = test_req
    else:
        requirement_text = args.requirement

    output = main(
        requirement=requirement_text,
        output_path=args.output,
        build_kb=not args.no_build_kb,
        model=args.model,
        max_chapters=args.max_chapters,
    )
    print(f"\n输出文件: {output}")
