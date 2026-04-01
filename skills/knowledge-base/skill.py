#!/usr/bin/env python3
"""
知识库Skill工具
支持多种文件格式 + 图片提取

Version: 1.1.0
Changelog:
  1.0.0 (2026-03-04) - 初始版本，支持本地LanceDB向量库
  1.1.0 (2026-03-30) - 新增docx图片提取功能
"""

import sys
import json
import os
import zipfile
import io
import requests
import pyarrow as pa
from pathlib import Path
from datetime import datetime
import hashlib
import shutil

DATA_DIR = Path.home() / "openclaw_doc" / "knowledge"
DB_DIR = DATA_DIR / "lancedb"
DOCS_DIR = DATA_DIR / "docs"
IMAGES_DIR = DATA_DIR / "images"
OLLAMA_URL = "http://localhost:11434/api/embeddings"
VECTOR_DIM = 768

# 全局缓存
_table = None

def get_embedding(text: str) -> list[float]:
    """获取文本向量"""
    resp = requests.post(OLLAMA_URL, json={"model": "nomic-embed-text", "prompt": text}, timeout=60)
    emb = resp.json()["embedding"]
    return emb[:VECTOR_DIM] + [0.0] * (VECTOR_DIM - len(emb))

def get_table():
    """获取向量表"""
    global _table
    if _table is None:
        import lancedb
        db = lancedb.connect(str(DB_DIR))
        try:
            _table = db.open_table("documents")
        except:
            schema = pa.schema([
                ("id", pa.string()), ("title", pa.string()), ("content", pa.string()),
                ("keywords", pa.list_(pa.string())), ("source_type", pa.string()),
                ("created", pa.string()), ("vector", pa.list_(pa.float32(), VECTOR_DIM)),
                ("images", pa.list_(pa.string())),  # 新增图片路径字段
            ])
            _table = db.create_table("documents", schema=schema)
    return _table

def extract_images_from_docx(file_path: str, doc_id: str) -> list:
    """从Word文档中提取图片"""
    images = []
    IMAGES_DIR.mkdir(parents=True, exist_ok=True)
    
    try:
        with zipfile.ZipFile(file_path, 'r') as zip_ref:
            # 查找word/media目录下的所有图片
            for name in zip_ref.namelist():
                if name.startswith('word/media/') and not name.endswith('/'):
                    # 获取图片扩展名
                    ext = os.path.splitext(name)[1].lower()
                    if ext in ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp']:
                        # 读取图片数据
                        img_data = zip_ref.read(name)
                        
                        # 生成唯一文件名
                        img_name = f"{doc_id}_{len(images)}{ext}"
                        img_path = IMAGES_DIR / img_name
                        
                        # 保存图片
                        with open(img_path, 'wb') as f:
                            f.write(img_data)
                        
                        images.append(str(img_path))
                        print(f"  📷 提取图片: {img_name}")
    except Exception as e:
        print(f"  ⚠️ 图片提取失败: {e}")
    
    return images

def extract_text_from_file(file_path: str, extract_images: bool = True) -> tuple:
    """从各种文件中提取文本和图片"""
    ext = Path(file_path).suffix.lower()
    images = []
    
    # 纯文本
    if ext in ['.txt', '.md', '.markdown']:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read(), []
    
    # Word文档
    elif ext in ['.docx', '.doc']:
        try:
            from docx import Document
            doc = Document(file_path)
            
            # 提取文本
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # 提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        text_parts.append(row_text)
            
            content = '\n'.join(text_parts)
            
            # 提取图片
            if extract_images:
                doc_id = hashlib.md5(file_path.encode()).hexdigest()[:8]
                images = extract_images_from_docx(file_path, doc_id)
            
            return content, images
        except Exception as e:
            return f"Word解析失败: {e}", []
    
    # PDF
    elif ext == '.pdf':
        try:
            from pypdf import PdfReader
            reader = PdfReader(file_path)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            return text, []
        except:
            return "", []
    
    # Excel
    elif ext in ['.xlsx', '.xls']:
        try:
            from openpyxl import load_workbook
            wb = load_workbook(file_path, read_only=True)
            text = ""
            for sheet in wb.sheetnames:
                ws = wb[sheet]
                for row in ws.iter_rows(values_only=True):
                    row_text = ' | '.join([str(cell) if cell is not None else '' for cell in row])
                    if row_text.strip():
                        text += row_text + "\n"
            return text, []
        except:
            return "", []
    
    # 未知格式
    else:
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read(), []
        except:
            return f"无法解析: {ext}", []

def add_doc(title: str, content: str = "", keywords=None, source_type="text", file_path: str = None, images: list = None) -> str:
    """添加文档"""
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    doc_id = hashlib.md5(f"{title}{datetime.now().isoformat()}".encode()).hexdigest()[:12]
    
    # 如果提供了文件路径，提取内容和图片
    if file_path and not content:
        content, extracted_images = extract_text_from_file(file_path)
        images = extracted_images
    
    # 保存原文
    (DOCS_DIR / f"{doc_id}.txt").write_text(content, encoding="utf-8")
    
    # 向量化存储
    emb = get_embedding(f"{title}\n{content[:1000]}")
    get_table().add([{
        "id": doc_id, 
        "title": title, 
        "content": content,
        "keywords": keywords or [], 
        "source_type": source_type,
        "created": datetime.now().isoformat(), 
        "vector": emb,
        "images": images or []
    }])
    return doc_id

def search(query: str, top_k: int = 3) -> list:
    """
    语义搜索 - 远程优先，本地兜底

    流程：
    1. 先搜索远程知识库
    2. 远程有结果 → 返回远程结果
    3. 远程无结果 → 搜索本地LanceDB
    4. 本地有结果 → 返回本地结果
    5. 都没有 → 返回空列表
    """
    # 1. 尝试远程知识库
    try:
        from knowledge_mcp import search as remote_search
        remote_results = remote_search(query, top_k=top_k)
        if isinstance(remote_results, dict) and remote_results.get('results'):
            results = remote_results['results']
            # 标记来源
            for r in results:
                r['_source'] = 'remote'
            print(f"[知识库] 远程搜索到 {len(results)} 条结果")
            return results
    except Exception as e:
        print(f"[知识库] 远程搜索失败: {e}，尝试本地...")

    # 2. 远程没有，搜索本地
    try:
        emb = get_embedding(query)
        local_results = get_table().search(emb, vector_column_name="vector").limit(top_k).to_list()
        if local_results:
            for r in local_results:
                r['_source'] = 'local'
            print(f"[知识库] 本地搜索到 {len(local_results)} 条结果")
            return local_results
    except Exception as e:
        print(f"[知识库] 本地搜索失败: {e}")

    # 3. 都没有
    print(f"[知识库] 远程和本地都没有搜索结果")
    return []

def format_result(r: dict) -> dict:
    """格式化搜索结果"""
    return {
        "title": r.get("title", ""),
        "overview": r.get("content", "")[:200],
        "content": r.get("content", ""),
        "images": r.get("images", []),
        "has_images": len(r.get("images", [])) > 0
    }

if __name__ == "__main__":
    cmd = sys.argv[1] if len(sys.argv) > 1 else "search"
    
    if cmd == "add":
        title = sys.argv[2] if len(sys.argv) > 2 else "无标题"
        file_path = sys.argv[3] if len(sys.argv) > 3 else None
        keywords = sys.argv[4].split(",") if len(sys.argv) > 4 else []
        
        content, images = extract_text_from_file(file_path) if file_path else ("", [])
        doc_id = add_doc(title, content, keywords, file_path=file_path, images=images)
        print(f"✅ 已保存 (ID: {doc_id}, 图片: {len(images)}张)")
    
    elif cmd == "search":
        query = sys.argv[2] if len(sys.argv) > 2 else "test"
        results = search(query)
        print(json.dumps([format_result(r) for r in results], ensure_ascii=False, indent=2))
    
    elif cmd == "extract":
        file_path = sys.argv[2] if len(sys.argv) > 2 else ""
        content, images = extract_text_from_file(file_path)
        print(f"提取到 {len(content)} 字符, {len(images)} 张图片")
        for img in images:
            print(f"  📷 {img}")
    
    else:
        print("Usage: python skill.py [add|search|extract] [args...]")
